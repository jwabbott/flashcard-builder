# # Flashcard Builder
# ## Joshua W. Abbott
# ### Description:
# This code creates printable flashcard sheets for learning student names and faces.
# It accepts as input a CSV file containing student information, including thair name and some basic biographical info, along with image files containing the students' headshots.
# It creates a Word file containing cards for 10 students on each sheet, which can be printed, cut out, and used as flashcards.

# ## Install tools

import pandas as pd
import numpy as np
import math
import argparse
from pathlib import Path
from PIL import Image
import os
import tempfile
import shutil
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE


# Set values for page and card size, orientation, and margins (all measurements in inches unless otherwise indicated)

PAGE_HEIGHT = 8.5
PAGE_WIDTH = 11
PAGE_MARGINS = 0.25
ROWS = 2
COLUMNS = 5
CARDS_PER_PAGE = ROWS * COLUMNS
TABLE_HEIGHT = 0.9 * PAGE_HEIGHT
TABLE_WIDTH = PAGE_WIDTH
CARD_HEIGHT = TABLE_HEIGHT / ROWS
CARD_WIDTH = TABLE_WIDTH / COLUMNS
IMG_MAX_HEIGHT = CARD_HEIGHT * 0.9
IMG_MAX_WIDTH = CARD_WIDTH * 0.9
NAME_FONT_PT = 14
INFO_FONT_PT = 9


def build_image_index(image_folder):

	"""
	Goes through a specified folder and creates an index dictionary of the filenames (key: filename, value: full path).
	"""

	image_dict = {}

	for file_name in os.listdir(image_folder):					# Iterate through all items in the image folder directory, and create the full path for each file

		full_path = os.path.join(image_folder, file_name)
		if os.path.isfile(full_path):						# if the item is a file (i.e. not a folder), add its filename to the dictionary
			image_dict[file_name] = full_path
	            
	return image_dict


def format_image (image, filename, temp_image_folder):

	"""
	opens and formats the given image, then saves formatted copy of image for insertion into Word document
	then stores formatting details in a dictionary that can be added to corresponding row in the main dataframe
	"""

	image_formatting_info = {}
	image_formatting_info['image_path'] = image

	with Image.open(image) as img:

		img_size_tuple = img.size
		img_width_px, img_height_px = img_size_tuple
		image_formatting_info['width_px'] = img_width_px			# add image dimensions in pixels to the dictionary
		image_formatting_info['height_px'] = img_height_px

		dpi_val = img.info.get('dpi')						# get the "dots-per-inch" from the image info and add to dictionary
		if dpi_val:								# use 72 as the default dpi if not contained in image info or if <=0
			candidate_dpi = dpi_val[0]
			if candidate_dpi > 0 and candidate_dpi < 100:
				safe_dpi = candidate_dpi
			else: safe_dpi = 72
		else:
			safe_dpi = 72

		image_formatting_info['dpi'] = safe_dpi

		img_height = img_height_px / safe_dpi					# calculate the image dimensions in inches
		img_width = img_width_px / safe_dpi
		image_formatting_info['height_in'] = img_height
		image_formatting_info['width_in'] = img_width

		scale_h = IMG_MAX_HEIGHT / img_height					# calculate the scaling factor to fit the image in the table cell
		scale_w = IMG_MAX_WIDTH / img_width
		scale = min(scale_w, scale_h, 1.0)
		image_formatting_info['scale'] = scale
	
		scaled_height = img_height * scale					# calculate scaled image dimensions in inches
		scaled_width = img_width * scale
		image_formatting_info['scaled_height_in'] = scaled_height
		image_formatting_info['scaled_width_in'] = scaled_width

		scaled_height_px = int(img_height_px * scale)				# calculate scaled image dimensions in pixels
		scaled_width_px = int(img_width_px * scale)
		image_formatting_info['scaled_height_px'] = scaled_height_px
		image_formatting_info['scaled_width_px'] = scaled_width_px

		# save copy of image in a temporary file using a safe value for dpi
		temp_image_filename = 'temp_'+filename
		temp_image_path = os.path.join(temp_image_folder, temp_image_filename)	
		img.save(temp_image_path, dpi=(safe_dpi, safe_dpi))
		image_formatting_info['temp_path'] = temp_image_path


	return image_formatting_info							# return dictionary with all image info


def add_front(document, batch_df):

	"""
	Adds front table with student info for current batch to Word document.
	"""

	table = document.add_table(rows=2, cols=5)					# insert table with fixed row and column sizes
	table.style = 'Table Grid' 
	table.allow_autofit = False

	for column in table.columns:
		column.width = Inches(CARD_WIDTH)
		for cell in column.cells:
			cell.width = Inches(CARD_WIDTH)

	for row in table.rows:
		row.height = Inches(CARD_HEIGHT)
		row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
											

	student_index = 0

	for row in table.rows:								# iterate through each cell in the table
		for cell in row.cells:

			if student_index < len(batch_df):				# only add info to cell is there are more students in the batch

				# Center text, set to 14-point font size, enter student name

				para = cell.add_paragraph()
				para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
				name_text = batch_df.iat[student_index, 0]
				run = para.add_run(name_text + "\n")
				run.bold = True	
				run.font.size = Pt(14)
		

				# Left justify text, reset to 11-pt font size, enter other student info
		
				new_para = cell.add_paragraph()
				new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
				new_para_format = new_para.paragraph_format
				new_para_format.left_indent = Inches(0.25)


				# iterate through student info and add to card

				for i in range(3):
					line = i + 1
					line_text = batch_df.iat[student_index, line]
					run = new_para.add_run(line_text + "\n")
					run.bold = False	
					run.font.size = Pt(11)


			student_index = student_index + 1				# iterate to next student in the batch


def add_back(document, batch_df):

	"""
	Adds back table with student photos to Word document.
	"""

	table = document.add_table(rows=2, cols=5)					# insert table with fixed row and column sizes
	table.style = 'Table Grid' 
	table.allow_autofit = False

	for column in table.columns:
		column.width = Inches(CARD_WIDTH)
		for cell in column.cells:
			cell.width = Inches(CARD_WIDTH)

	for row in table.rows:
		row.height = Inches(CARD_HEIGHT)
		row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


	student_index = 0


	for row in table.rows:								# iterate through each cell in the table, 
		for cell in reversed(row.cells):					# but in such a way as to pair each photo on the back to the matching student info on the front

			if student_index < len(batch_df):				# only add photos if there are more in the batch, up to 10

				para = cell.add_paragraph()
				para.alignment = WD_ALIGN_PARAGRAPH.CENTER
				run = para.add_run()

				photo_path = batch_df.loc[student_index, 'temp_path']
				photo_width = batch_df.loc[student_index, 'scaled_width_in']

				run.add_picture(photo_path, width=Inches(photo_width))
				cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

			student_index = student_index + 1				# iterate to the next studdent in the batch


def create_doc(info_path, images_folder, output_doc):

	"""
	Creates a new MS Word document (in .docx format) and enters the student information (for 10 students at a time) onto the front of each page 
	and the corresponding headshot images onto the back of each page.
	"""

	df = pd.read_csv(info_path)							# Read the student information from the csv file into a dataframe and display the first few rows
	df.fillna('', inplace=True)
	print (df.head())

	image_dict = build_image_index(images_folder)					# Create an index dictionary of the headshot image filenames and add their filepaths to the dataframe
	df["image_path"] = df["image_filename"].map(image_dict)

	temp_image_folder = "temp_image_folder"						# Create a folder to hold temporary copies of headshot image files
	os.makedirs(temp_image_folder, exist_ok=True)

											# Create a new dataframe for image formatting data

	image_formatting_columns = ['image_path', 'height_px', 'width_px', 'height_in', 'width_in', 'scale', 'scaled_height_in', 'scaled_width_in', 'scaled_height_px', 'scaled_width_px', 'temp_path']
	image_info_df = pd.DataFrame(columns=image_formatting_columns)
	for row in range(len(df)):
		path_of_image_to_format = df.at[row, 'image_path']
		filename_of_image_to_format = df.at[row, 'image_filename']
		image_info_df.loc[row] = format_image(path_of_image_to_format, filename_of_image_to_format, temp_image_folder)


	df = pd.merge(df, image_info_df, on='image_path')				# Merge the image formatting dataframe with the student information dataframe

	
	document = Document()								# Create a new blank Word document

	section = document.sections[-1]							# Set the document's orientation to landscape
	new_width, new_height = section.page_height, section.page_width
	section.page_width = new_width
	section.page_height = new_height

	section.top_margin = Inches(PAGE_MARGINS)					# Set all the document's margins to a quarter inch
	section.bottom_margin = Inches(PAGE_MARGINS)
	section.left_margin = Inches(PAGE_MARGINS)
	section.right_margin = Inches(PAGE_MARGINS)

	num_of_batches = math.ceil(len(df) / 10)					# calculate the number of batches of 10 students each, with last batch containing any remaining

	for batch_index in range(num_of_batches):					# iterate through each batch of students

		start = batch_index * 10
		end = start + 10
		end = min(end, len(df))

		batch_df = df.iloc[start:end].reset_index(drop=True)			# create a subset dataframe containing only the info for the students in the current batch

		add_front(document, batch_df)						# insert student info and photos into Word document by calling the functions to add front and back

		add_back(document, batch_df)


	document.save('flashcards.docx')						# save the Word document
	shutil.rmtree(temp_image_folder)						# delete temporary images


def parse_args():

	p = argparse.ArgumentParser()
	p.add_argument("info", help="full path of csv file containing student info")
	p.add_argument("images", help="name of folder containing headshot images")
	p.add_argument("output", help="filename of output Word document containing flashcards: __.docx")
	return p.parse_args()


def main():

	args = parse_args()

	info_path = Path(args.info)
	images_folder = Path(args.images)
	output_doc = Path(args.output)

	if not info_path.exists():
		logger.error("CSV file with student information not found: %s", info_path)
		sys.exit(2)
	if not images_folder.exists() or not image_folder.is_dir():
		logger.error("Folder with headshot images not found or not a directory: %s", images_folder)
		sys.exit(2)

	create_doc(info_path, images_folder, output_doc)
	


if __name__ == "__main__":
    main()
